import os
import re
import docx
import PyPDF2
import string
from flask import Flask, request, jsonify
from flask_cors import CORS
from spacy.lang.en import English
from spacy.lang.en.stop_words import STOP_WORDS
from dotenv import load_dotenv
import google.generativeai as genai
import spacy # Import spacy directly for loading models

# NEW IMPORTS FOR RAG (Retrieval-Augmented Generation)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)
CORS(app)

# ----- Configuration -----
MAX_WORD_LIMIT = 5000# Maximum words for processing a single section/chunk
CHUNK_THRESHOLD = 10000# Use chunk summarization when document exceeds this many words
FINAL_NUM_SENTENCES = 12 # For full-document summarization (for PDFs)
CHUNK_SIZE = 5000# Chunk size for long documents

# ----- Minimal NLP Pipeline (for sentence segmentation) -----
def get_minimal_nlp():
    nlp_min = English()
    nlp_min.add_pipe("sentencizer")
    return nlp_min

nlp_min = get_minimal_nlp()

# ----- SpaCy Model for Chatbot (download if not present) -----
try:
    nlp_chat = spacy.load("en_core_web_sm")
except OSError:
    print("Downloading en_core_web_sm model for chatbot. This may take a moment...")
    spacy.cli.download("en_core_web_sm")
    nlp_chat = spacy.load("en_core_web_sm")

# ----- Google Gemini API Configuration -----
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY not found in environment variables. Please set it in your .env file.")

genai.configure(api_key=GEMINI_API_KEY)

# Initialize the generative model with the recommended model
gemini_model = genai.GenerativeModel('gemini-1.5-flash')

# ----- GLOBAL VARIABLES FOR RAG & METADATA -----
# This will store our FAISS vector index (the searchable database of document chunks)
vector_store = None
# This will store the complete raw text of the last uploaded document
full_document_text = ""
# This still stores the last generated summary for initial display on the frontend
last_generated_summary = ""
# Global variable to store extracted metadata like authors, title
document_metadata = {
    "authors": [],
    "title": ""
}

# ----- Enhanced Cleaning Function -----
def clean_text(text):
    """
    Cleans the input text by removing extraneous academic front-matter,
    such as DOI lines, email addresses, bibliographic details, and other
    clutter that is usually not desired in a clean summary.
    """
    # Remove email addresses
    text = re.sub(r'\S+@\S+', ' ', text)
    # Remove DOI lines or similar patterns
    text = re.sub(r'Digital\s*Object\s*Identifier.*', ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'DOI[:\s]+[^\s]+', ' ', text, flags=re.IGNORECASE)
    # Remove corresponding author lines
    text = re.sub(r'Corresponding author:.*', ' ', text, flags=re.IGNORECASE)
    # Remove citation markers in brackets, e.g., [45]
    text = re.sub(r'\[\s*[0-9]+(?:\s*,\s*[0-9]+)\s\]', ' ', text)
    # Remove standalone numeric blocks that often come from affiliations
    text = re.sub(r'\b\d{2,}\b', ' ', text)
    # If the document is academic, keep only text starting at the ABSTRACT marker.
    if "abstract" in text.lower():
        parts = re.split(r'abstract', text, flags=re.IGNORECASE)
        if len(parts) > 1:
            text = "ABSTRACT " + parts[1]
    # Remove extra white spaces and stray punctuation
    text = re.sub(r'\s+', ' ', text).strip()
    return text

# ----- Post-Processing Refinement Layer -----
def refine_summary(summary):
    """
    Applies a sentence-level cleanup on the generated summary:
      - Removes dangling phrases like "INDEX TERMS".
      - Fixes spacing around punctuation.
      - Appends terminal punctuation (period) to sentences that may have incomplete phrases.
      - Performs basic grammar and spacing fixes.
    """
    # Remove unwanted dangling phrases, e.g., "INDEX TERMS"
    summary = re.sub(r'\bINDEX\s+TERMS\b', '', summary, flags=re.IGNORECASE)
    
    # Fix spacing: remove extra spaces before punctuation
    summary = re.sub(r'\s+([,.!?;:])', r'\1', summary)
    
    # Split the summary into sentences (assuming punctuation marks end sentences)
    sentences = re.split(r'(?<=[.!?])\s+', summary)
    refined_sentences = []
    for sent in sentences:
        s = sent.strip()
        # If a sentence doesn't end with terminal punctuation, add one.
        if s and s[-1] not in ".!?":
            s = s + "."
        refined_sentences.append(s)
    refined_summary = " ".join(refined_sentences)
    
    # Remove extra spaces
    refined_summary = re.sub(r'\s+', ' ', refined_summary).strip()
    return refined_summary

# ----- Helper Function: Heading Detector -----
def is_heading_sentence(sentence):
    """
    A simple heuristic: if a sentence is short (e.g., 3-6 words)
    and is written in all uppercase, it is likely a heading.
    """
    text = sentence.text.strip()
    words = text.split()
    if 3 <= len(words) <= 6 and text.isupper():
        return True
    return False

# ----- Extract Text from PDF or DOCX (ENHANCED VERSION) -----
def extract_text(file):
    text = ""
    filename = file.filename.lower()
    # Reset metadata for each new file
    global document_metadata # Ensure we can modify the global variable
    document_metadata["authors"] = []
    document_metadata["title"] = ""

    if filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        # Attempt to get title and author from PDF metadata
        if reader.metadata:
            if reader.metadata.get('/Title'):
                document_metadata["title"] = reader.metadata['/Title'].strip()
            if reader.metadata.get('/Author'):
                author_string = reader.metadata['/Author'].strip()
                document_metadata["authors"] = [a.strip() for a in author_string.split(';')] if ';' in author_string else [author_string.strip()]

        full_pdf_text_list = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                full_pdf_text_list.append(page_text)
                
                # For PDF, try to find authors/title from first few pages' content
                if i < 3: # Check first 3 pages
                    lines = page_text.split('\n')
                    for line in lines[:15]: # Check top few lines of these pages
                        line_stripped = line.strip()

                        # Improved regex for authors (handles ALL CAPS, numbers, and common separators)
                        # Looks for sequences of capitalized words/initials/numbers, separated by commas, 'and', or spaces
                        author_match = re.search(r'([A-Z][A-Z0-9\s.,-]+(?:(?: and|AND)\s+)?)+', line_stripped)
                        if author_match and len(line_stripped.split()) < 20 and not line_stripped.startswith("Digital Object Identifier"): # Heuristic to avoid DOIs/headers
                            potential_authors_str = author_match.group(0).strip()
                            # Clean up potential author string: remove leading/trailing numbers/spaces/commas
                            potential_authors_str = re.sub(r'^\d+\s*|,\s*\d+$', '', potential_authors_str).strip()
                            potential_authors_str = re.sub(r'\s*\d+\s*$', '', potential_authors_str).strip() # Remove trailing numbers
                            potential_authors_str = re.sub(r'\s*,\s*', ';', potential_authors_str) # Standardize separator for splitting
                            
                            # Split by common separators (;, AND) and clean up individual names
                            names = re.split(r';|AND\s*', potential_authors_str, flags=re.IGNORECASE)
                            cleaned_names = []
                            for name in names:
                                cleaned_name = name.strip()
                                # Remove any remaining single numbers or trailing numbers from names
                                cleaned_name = re.sub(r'\s*\b\d+\b\s*', '', cleaned_name).strip()
                                if cleaned_name and len(cleaned_name.split()) >= 2: # Ensure it looks like a name
                                    cleaned_names.append(cleaned_name)
                            
                            if cleaned_names and not document_metadata["authors"]: # If authors not already found
                                document_metadata["authors"].extend(cleaned_names)
                                # Remove duplicates in case multiple heuristics find them
                                document_metadata["authors"] = list(dict.fromkeys(document_metadata["authors"])) # Preserve order, remove duplicates
                                print(f"DEBUG: PDF Author (content) found: {document_metadata['authors']}") # Debug print

                        # Simple heuristic for title from first page
                        if not document_metadata["title"] and i == 0: # Only scan first page for title
                            # Look for a prominent line that appears to be a title
                            if len(line_stripped.split()) < 15 and len(line_stripped) > 20 and line_stripped.isupper():
                                document_metadata["title"] = line_stripped.strip()
                                print(f"DEBUG: PDF Title (content) found: {document_metadata['title']}") # Debug print

        text = "\n".join(full_pdf_text_list) # Combine all page text

    elif filename.endswith(".docx"):
        doc = docx.Document(file)
        full_docx_text_list = []
        # Attempt to get title and author from DOCX properties
        if doc.core_properties:
            if doc.core_properties.title:
                document_metadata["title"] = doc.core_properties.title.strip()
            if doc.core_properties.author:
                author_string = doc.core_properties.author.strip()
                document_metadata["authors"] = [a.strip() for a in author_string.split(';')] if ';' in author_string else [author_string.strip()]

        # Scan first few paragraphs for authors if not found in metadata
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if para_text:
                full_docx_text_list.append(para_text)
                if i < 5 and not document_metadata["authors"]: # Check first 5 paragraphs
                    # More robust author pattern detection for DOCX (can be similar to PDF or more structured)
                    # This regex tries to capture multiple names separated by commas, "and", or "by"
                    author_match = re.search(r'(?:by|authors?|contributed\s+by)[:\s]*([A-Za-z0-9\s.,-]+(?:(?: and|AND)\s+)?)+', para_text, re.IGNORECASE)
                    if author_match:
                        potential_authors_str = author_match.group(1).strip() # Group 1 is the actual names
                        potential_authors_str = re.sub(r'\s*\d+\s*$', '', potential_authors_str).strip() # Remove trailing numbers
                        potential_authors_str = re.sub(r'\s*,\s*', ';', potential_authors_str) # Standardize separator
                        
                        names = re.split(r';|AND\s*', potential_authors_str, flags=re.IGNORECASE)
                        cleaned_names = []
                        for name in names:
                            cleaned_name = name.strip()
                            cleaned_name = re.sub(r'\s*\b\d+\b\s*', '', cleaned_name).strip() # Remove any remaining single numbers
                            if cleaned_name and len(cleaned_name.split()) >= 2:
                                cleaned_names.append(cleaned_name)

                        if cleaned_names:
                            document_metadata["authors"].extend(cleaned_names)
                            document_metadata["authors"] = list(dict.fromkeys(document_metadata["authors"]))
                            print(f"DEBUG: DOCX Author (content) found: {document_metadata['authors']}") # Debug print

        text = " ".join(full_docx_text_list)
    else:
        raise ValueError("Only PDF and DOCX files are supported.")
    
    print(f"DEBUG: Final extracted authors: {document_metadata['authors']}")
    print(f"DEBUG: Final extracted title: {document_metadata['title']}")
    return text

# ----- Extract Structured Text from DOCX (Kept if needed for section-wise summarization display, NOT for RAG ingestion) -----
def extract_structured_text_docx(file):
    """
    Extracts sections from a DOCX file using paragraph styles.
    A paragraph is considered a heading if its style name (case-insensitive)
    contains "heading". Subsequent paragraphs are grouped under that heading.
    """
    # IMPORTANT: For RAG ingestion (creating vector_store), the 'extract_text' function
    # is now preferred as it extracts full text and metadata regardless of structure.
    # This function is mainly for generating section-wise summaries if desired for display.
    
    doc = docx.Document(file)
    sections = []
    current_heading = None
    current_section_text = ""
    for para in doc.paragraphs:
        style_name = para.style.name.lower() if para.style and para.style.name else ""
        if "heading" in style_name:
            if current_heading is not None:
                sections.append((current_heading, current_section_text.strip()))
            current_heading = para.text.strip()
            current_section_text = ""
        else:
            if para.text.strip():
                current_section_text += " " + para.text.strip()
    if current_heading is not None:
        sections.append((current_heading, current_section_text.strip()))
    else: # If no headings found, treat entire doc as one section
        full_text = " ".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
        sections = [("Document Content", full_text)] # Provide a default heading
    return sections

# ----- Limit Text -----
def limit_text(text, word_limit=MAX_WORD_LIMIT):
    words = text.split()
    if len(words) > word_limit and len(words) < CHUNK_THRESHOLD:
        return " ".join(words[:word_limit])
    return text

# ----- Enhanced Frequency-based Summarization -----
def summarize_text_spacy(text, num_sentences=FINAL_NUM_SENTENCES):
    """
    Summarizes text by:
      - First cleaning the text to remove extraneous bibliographic info.
      - Segmenting the text into sentences using the minimal SpaCy pipeline.
      - Building a word frequency table from alphabetic tokens (ignoring punctuation, numbers, and stopwords).
      - Scoring each sentence (adding bonus points if the sentence looks like a heading).
      - Selecting and returning the top scored sentences ordered by their original sequence.
    """
    text = clean_text(text)
    doc = nlp_min(text)
    sentences = list(doc.sents)
    
    freq = {}
    for word in text.lower().split():
        word_cleaned = word.strip(string.punctuation)
        if not word_cleaned or not word_cleaned.isalpha() or word_cleaned in STOP_WORDS:
            continue
        freq[word_cleaned] = freq.get(word_cleaned, 0) + 1

    sentence_scores = {}
    for sent in sentences:
        words = [w.strip(string.punctuation) for w in sent.text.lower().split()]
        clean_words = [w for w in words if w and w.isalpha() and w not in STOP_WORDS]
        score = sum(freq.get(word, 0) for word in clean_words)
        # Bonus points for heading-like sentences.
        if is_heading_sentence(sent):
            score += 50
        sentence_scores[sent] = score

    # Select top sentences. Ensure num_sentences does not exceed available sentences.
    actual_num_sentences = min(num_sentences, len(sentences))
    ranked = sorted(sentence_scores, key=sentence_scores.get, reverse=True)[:actual_num_sentences]
    ranked.sort(key=lambda s: s.start) # Sort by original order
    summary = " ".join(sent.text.strip() for sent in ranked)
    return summary

# ----- Chunk Summarization for Very Long Documents -----
def chunk_summarize(text, num_sentences=FINAL_NUM_SENTENCES, chunk_size=CHUNK_SIZE):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i+chunk_size])
        chunks.append(chunk)

    chunk_summaries = []
    # Ensure at least 1 sentence per chunk if num_sentences is very low or chunks are many
    per_chunk_sentences = max(1, num_sentences // len(chunks))
    for chunk in chunks:
        chunk_summary = summarize_text_spacy(chunk, num_sentences=per_chunk_sentences)
        chunk_summaries.append(chunk_summary)

    combined_summary_text = " ".join(chunk_summaries)
    final_summary = summarize_text_spacy(combined_summary_text, num_sentences=num_sentences)
    return final_summary

# ----- Dispatcher Summarization for Full Document (PDF) -----
def summarize_document(text, num_sentences=FINAL_NUM_SENTENCES):
    text = clean_text(text)
    words = text.split()
    if len(words) > CHUNK_THRESHOLD:
        return chunk_summarize(text, num_sentences=num_sentences)
    else:
        limited_text = limit_text(text)
        return summarize_text_spacy(limited_text, num_sentences=num_sentences)

# ----- Section-wise Summarization with Headings for DOCX -----
def summarize_sections(sections, sentences_per_section=2):
    """
    Summarizes each section obtained from a DOCX file.
    Each section is cleaned and summarized individually while preserving its heading.
    """
    summaries = []
    for heading, content in sections:
        content = clean_text(content)
        # Only summarize if content is substantial
        if not content.strip() or len(content.split()) < 30:
            section_summary = content.strip()
        else:
            sect_text = limit_text(content, word_limit=MAX_WORD_LIMIT)
            section_summary = summarize_text_spacy(sect_text, num_sentences=sentences_per_section)
        
        if heading and heading != "Document Content":
            summaries.append(f"\n### {heading} ###\n{section_summary}\n")
        else:
            summaries.append(section_summary)
    return "\n".join(summaries)

# ----- Flask Upload Route (UPDATED FOR RAG INGESTION) -----
@app.route("/upload", methods=["POST"])
def upload():
    global vector_store
    global full_document_text
    global last_generated_summary # Still update this for the initial summary display
    global document_metadata # Make sure to access global metadata

    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    try:
        filename = file.filename.lower()
        if not (filename.endswith(".pdf") or filename.endswith(".docx")):
            return jsonify({"error": "Only PDF and DOCX files are supported."}), 400

        # --- STEP 1: Extract Full Text & Metadata ---
        # The enhanced extract_text function now handles populating document_metadata
        # and returning the full text regardless of file type.
        full_document_text = extract_text(file) # This also populates document_metadata internally

        if not full_document_text.strip():
            return jsonify({"error": "No readable text found in the document."}), 400

        # --- STEP 2: RAG INGESTION PROCESS (Text Splitting, Embedding, Vector Store Creation) ---
        print("Starting RAG ingestion: splitting text and creating embeddings...")

        # 2.1 Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        texts = text_splitter.split_text(full_document_text)
        print(f"Document split into {len(texts)} chunks.")

        # 2.2 Create embeddings model (this model runs locally, downloading if first time)
        embeddings_model = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
        print("Embedding model loaded.")

        # 2.3 Create a FAISS vector store from the chunks and embeddings
        vector_store = FAISS.from_texts(texts, embeddings_model)
        print("Vector store created successfully.")

        # --- STEP 3: Generate Initial Summary (for display, not chatbot's primary Q&A) ---
        initial_summary = summarize_document(full_document_text, num_sentences=FINAL_NUM_SENTENCES)
        refined_summary = refine_summary(initial_summary)
        last_generated_summary = refined_summary # Store for initial chatbot message/display

        return jsonify({"summary": refined_summary, "message": "Document uploaded and indexed for Q&A. You can now ask questions directly about the document."})
    except Exception as e:
        print(f"Error during file upload and RAG ingestion: {e}")
        vector_store = None # Clear vector store on error
        full_document_text = ""
        last_generated_summary = "" # Clear summary on error too
        # Also clear metadata on error
        document_metadata["authors"] = []
        document_metadata["title"] = ""
        return jsonify({"error": str(e)}), 500

# ----- Chatbot Route (UPDATED FOR RAG QUERIES AND METADATA CHECK) -----
@app.route("/chatbot", methods=["POST"])
def chatbot():
    user_message = request.json.get("message")
    if not user_message:
        return jsonify({"error": "No message provided"}), 400
    
    # Check if a document has been uploaded and indexed
    if vector_store is None:
        return jsonify({"response": "Please upload a document first so I can answer questions about it."}), 200

    # --- NEW: Check for specific metadata questions first ---
    # Using a more robust check for author/title related queries
    lower_user_message = user_message.lower()
    if "author" in lower_user_message or "who wrote" in lower_user_message or "written by" in lower_user_message:
        if document_metadata["authors"]:
            return jsonify({"response": f"The author(s) of this document appear to be: {', '.join(document_metadata['authors'])}."})
        else:
            return jsonify({"response": "I couldn't find specific author information in the document's metadata or initial content scan."})

    if "title" in lower_user_message or "name of the document" in lower_user_message:
        if document_metadata["title"]:
            return jsonify({"response": "The title of this document is: " + document_metadata["title"]})
        else:
            return jsonify({"response": "I couldn't find specific title information in the document's metadata or initial content scan."})
    # --- END NEW METADATA CHECK ---

    try:
        # --- STEP 1: Retrieval ---
        print(f"Retrieving relevant chunks for query: '{user_message}'")
        retrieved_docs = vector_store.similarity_search(user_message, k=4) 
        
        context_from_document = "\n\n".join([doc.page_content for doc in retrieved_docs])
        print(f"Retrieved {len(retrieved_docs)} chunks. Context length: {len(context_from_document)} characters.")

        # Ensure there's context retrieved before asking Gemini
        if not context_from_document.strip():
            return jsonify({"response": "I cannot answer that question based on the provided document as I couldn't find relevant information."})

        # --- STEP 2: Augmentation & Generation ---
        prompt_template = PromptTemplate(
            template="""You are a helpful assistant specialized in answering questions based only on the provided document excerpts.
            If the answer cannot be found in the provided document excerpts, respond with "I cannot answer that question based on the provided document."
            Do not make up information.
            
            Document excerpts:
            {context}

            Question: {question}
            Answer:""",
            input_variables=["context", "question"],
        )

        final_prompt = prompt_template.format(context=context_from_document, question=user_message)
        print("Sending augmented prompt to Gemini model for generation...")
        
        response = gemini_model.generate_content(final_prompt)
        bot_response = response.text
        print("Gemini response received.")

    except Exception as e:
        print(f"Chatbot RAG process failed: {e}")
        bot_response = "I'm sorry, I encountered an internal error while trying to answer your question from the document. Please try again."
    
    return jsonify({"response": bot_response})

if __name__ == "__main__":
    # Ensure debug is set to True only during development
    app.run(debug=True)